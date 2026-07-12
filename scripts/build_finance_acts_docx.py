"""
Extract GST-related provisions from Finance Acts (2017 onward) and compile
into Finance_Acts_GST_Provisions.docx for portal ingestion.

Sources: www.indiabudget.gov.in (Finance Bills), gstcouncil.gov.in (enacted
GST extracts), egazette.gov.in (enacted Finance Acts).
"""
from __future__ import annotations

import os
import re
from datetime import datetime

import fitz
from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
SRC_DIR = os.path.join(ROOT, "data", "finance_acts_source")
OUT_DOCX = os.path.join(ROOT, "Finance_Acts_GST_Provisions.docx")

# Finance Act year -> primary PDF source (relative to SRC_DIR) and citation URL
ACT_SOURCES: list[dict] = [
    {
        "year": 2017,
        "title": "Finance Act, 2017",
        "note_only": True,
        "note": (
            "Finance Act, 2017 (assented 31 March 2017) predates the GST laws, which were "
            "enacted separately in April 2017 and brought into force from 1 July 2017. "
            "This Act contains no amendments to the CGST, IGST, UTGST or Compensation Cess Acts."
        ),
        "source_url": "https://egazette.gov.in/WriteReadData/2017/175141.pdf",
        "source_label": "eGazette — Finance Act, 2017",
    },
    {
        "year": 2018,
        "title": "Finance Act, 2018",
        "pdf": "egazette_fa2018.pdf",
        "fallback_pdf": "indiabudget_fb_2018-19.pdf",
        "source_url": "https://egazette.gov.in/writereaddata/2018/184302.pdf",
        "source_label": "eGazette — Finance Act, 2018",
        "budget_url": "https://www.indiabudget.gov.in/budget2018-2019/ub2018-19/fb/bill.pdf",
    },
    {
        "year": 2019,
        "title": "Finance Act, 2019",
        "note_only": True,
        "note": (
            "Finance Act, 2019 (assented 21 February 2019) contains no amendments to the "
            "CGST, IGST, UTGST or Compensation Cess Acts. GST amendments for 2019-20 were "
            "enacted through the Finance (No. 2) Act, 2019."
        ),
        "source_url": "https://gstcouncil.gov.in/sites/default/files/2024-04/finance_act_2019.pdf",
        "source_label": "GST Council — Finance Act, 2019",
    },
    {
        "year": 2019,
        "title": "Finance (No. 2) Act, 2019",
        "pdf": "indiabudget_fb_2019-20.pdf",
        "source_url": "https://www.indiabudget.gov.in/budget2019-20/doc/Finance_Bill.pdf",
        "source_label": "India Budget — Finance (No. 2) Bill, 2019",
    },
    {
        "year": 2020,
        "title": "Finance Act, 2020",
        "pdf": "gstcouncil_2020.pdf",
        "source_url": "https://gstcouncil.gov.in/sites/default/files/2024-04/finance_act_2020.pdf",
        "source_label": "GST Council — Finance Act, 2020",
    },
    {
        "year": 2021,
        "title": "Finance Act, 2021",
        "pdf": "gstcouncil_2021.pdf",
        "source_url": "https://gstcouncil.gov.in/sites/default/files/2024-04/finance_act_2021.pdf",
        "source_label": "GST Council — Finance Act, 2021",
    },
    {
        "year": 2022,
        "title": "Finance Act, 2022",
        "pdf": "gstcouncil_2022.pdf",
        "source_url": "https://gstcouncil.gov.in/sites/default/files/2024-04/finance_act_of_2022.pdf",
        "source_label": "GST Council — Finance Act, 2022",
    },
    {
        "year": 2023,
        "title": "Finance Act, 2023",
        "pdf": "gstcouncil_2023.pdf",
        "source_url": "https://gstcouncil.gov.in/sites/default/files/2024-04/finance_act_of_2023.pdf",
        "source_label": "GST Council — Finance Act, 2023",
    },
    {
        "year": 2024,
        "title": "Finance Act, 2024",
        "pdf": "fa2024_cf.pdf",
        "source_url": "https://egazette.gov.in/",
        "source_label": "eGazette — Finance Act, 2024 (Interim Budget, assented 15 February 2024)",
    },
    {
        "year": 2024,
        "title": "Finance (No. 2) Act, 2024",
        "pdf": "egazette_fa2024_no2.pdf",
        "fallback_pdf": "indiabudget_fb_2024-25.pdf",
        "source_url": "https://egazette.gov.in/WriteReadData/2024/256436.pdf",
        "source_label": "eGazette — Finance (No. 2) Act, 2024",
        "budget_url": "https://www.indiabudget.gov.in/budget2024-25/doc/Finance_Bill.pdf",
    },
    {
        "year": 2025,
        "title": "Finance Act, 2025",
        "pdf": "indiabudget_fb_2025-26.pdf",
        "source_url": "https://www.indiabudget.gov.in/budget2025-26/doc/Finance_Bill.pdf",
        "source_label": "India Budget — Finance Bill, 2025",
    },
]

GST_PART_START = re.compile(
    r"(?:PART|CHAPTER)\s+([IVXLC\d]+)\s*\n\s*"
    r"(AMENDMENTS?\s+TO\s+THE\s+)?"
    r"(CENTRAL\s+GOODS\s+AND\s+SERVICES\s+TAX\s+ACT,\s*2017|"
    r"INTEGRATED\s+GOODS\s+AND\s+SERVICES\s+TAX\s+ACT,\s*2017|"
    r"UNION\s+TERRITORY\s+GOODS\s+AND\s+SERVICES\s+TAX\s+ACT,\s*2017|"
    r"GOODS\s+AND\s+SERVICES\s+TAX\s+\(COMPENSATION\s+TO\s+STATES\)\s+ACT,\s*2017)",
    re.I,
)

CHAPTER_INDIRECT = re.compile(
    r"CHAPTER\s+IV\s*\n\s*INDIRECT\s+TAXES",
    re.I,
)

GST_SECTION_START = re.compile(
    r"\n(\d+)\.\s+"
    r"(?:In the|For section|After section|In section)\s+"
    r".{0,40}?"
    r"(?:Central Goods and Services Tax|Integrated Goods and Services Tax|"
    r"Union Territory Goods and Services Tax|Goods and Services Tax \(Compensation)",
    re.I | re.S,
)

GST_INLINE_START = re.compile(
    r"(?:^|\n)\s*Central\s+Goods\s+and\s+Services\s+Tax\s*\n\s*\d+\.",
    re.I,
)

GST_SUBHEAD_BODY = re.compile(
    r"\n\s*(Central Goods and Services Tax|Integrated Goods and Services Tax|"
    r"Union Territory Goods and Services Tax|"
    r"Goods and Services Tax \(Compensation[^\n]*)\s*\n\s*(\d+)\.\s+"
    r"(?:In the|For section|After section|In section)",
    re.I,
)

ENACTED_BODY = re.compile(r"BE it enacted by Parliament", re.I)

NON_GST_SUBHEAD = re.compile(
    r"\n\s*(Customs|Customs Duties|Service Tax|Central Excise|"
    r"CHAPTER\s+V)\s*\n",
    re.I,
)

END_MARKERS = [
    re.compile(r"\n\s*Service Tax\s*\n", re.I),
    re.compile(r"CHAPTER\s+V\s*\n\s*MISCELLANEOUS", re.I),
    re.compile(r"\n\s*THE\s+FIRST\s+SCHEDULE", re.I),
    re.compile(
        r"PART\s+I\s*\n\s*AMENDMENT TO THE UNIT TRUST OF INDIA",
        re.I,
    ),
]

NON_GST_PART = re.compile(
    r"(?:PART|CHAPTER)\s+[IVXLC\d]+\s*\n\s*AMENDMENTS?\s+TO\s+THE\s+(?!CENTRAL\s+GOODS|INTEGRATED\s+GOODS|UNION\s+TERRITORY|GOODS\s+AND\s+SERVICES\s+TAX)",
    re.I,
)

SCHEDULE_START = re.compile(
    r"\n\s*THE\s+(?:FIRST|SECOND|THIRD|FOURTH|FIFTH|SIXTH)\s+SCHEDULE",
    re.I,
)

SECTION_LINE = re.compile(
    r"^(\d+[A-Z]?)\.\s+(.+)$",
    re.M,
)

NOISE_LINE = re.compile(
    r"^(?:\(\w+\)|\d+\)|SEC\.\s*\d+|vlk/|EXTRAORDINARY|Hkkx|izkf/kdkj|PUBLISHED|"
    r"MINISTRY OF LAW|Legislative Department|New Delhi|Bill No\.|AS INTRODUCED|"
    r"ARRANGEMENT OF CLAUSES|CLAUSES$|\(viii\)$|\(ix\)$|\(x\)$)",
    re.I,
)


def pdf_text(path: str) -> str:
    doc = fitz.open(path)
    return "\n".join(doc.load_page(i).get_text() for i in range(doc.page_count))


def normalize_whitespace(text: str) -> str:
    text = text.replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def is_gst_only_extract(text: str) -> bool:
    upper = text.upper()
    if re.search(r"CHAPTER\s+II\s*\n\s*RATES", upper):
        return False
    if re.search(r"CHAPTER\s+III\s*\n\s*DIRECT\s+TAXES", upper):
        return False
    return text.lower().count("central goods and services tax") >= 2


def _find_block_end(text: str, start: int) -> int:
    end = len(text)
    for marker in END_MARKERS:
        m = marker.search(text, start + 20)
        if m:
            end = min(end, m.start())
    sched = SCHEDULE_START.search(text, start + 20)
    if sched:
        end = min(end, sched.start())
    nxt = NON_GST_PART.search(text, start + 20)
    if nxt:
        end = min(end, nxt.start())
    return end


def _enacted_text(text: str) -> str:
    match = ENACTED_BODY.search(text)
    return text[match.start() :] if match else text


def _extract_gst_subheading_blocks(text: str) -> list[tuple[str, str]]:
    """Extract CGST / IGST / UTGST / Compensation Cess blocks under indirect-tax headings."""
    enacted = _enacted_text(text)
    matches = list(GST_SUBHEAD_BODY.finditer(enacted))
    if not matches:
        return []

    blocks: list[tuple[str, str]] = []
    for i, match in enumerate(matches):
        heading = match.group(1).strip()
        start = match.start() + 1
        end = len(enacted)
        if i + 1 < len(matches):
            end = matches[i + 1].start()
        else:
            end = _find_block_end(enacted, start)
            non_gst = NON_GST_SUBHEAD.search(enacted, start + 40)
            if non_gst:
                end = min(end, non_gst.start())
        body = enacted[start:end].strip()
        if len(body) > 80:
            blocks.append((heading, body))
    return blocks


def _extract_gst_section_range(text: str) -> str:
    matches = list(GST_SECTION_START.finditer(text))
    if not matches:
        return ""

    body_matches = [
        m
        for m in matches
        if re.search(
            r"In the (?:Central|Integrated|Union Territory) Goods|"
            r"For section \d+ of the (?:Central|Integrated|Union Territory) Goods|"
            r"After section \d+ of the (?:Central|Integrated|Union Territory) Goods",
            text[m.start() : m.start() + 500],
            re.I,
        )
    ]
    if not body_matches:
        body_matches = matches[-5:]

    start = body_matches[0].start()
    end = _find_block_end(text, body_matches[-1].start())
    return text[start:end].strip()


def extract_gst_blocks(full_text: str) -> list[tuple[str, str]]:
    """Return list of (heading, body) for each GST-related part/chapter."""
    text = normalize_whitespace(full_text)

    if is_gst_only_extract(text):
        start = 0
        for marker in (
            "Central Goods and Services Tax",
            "CENTRAL GOODS AND SERVICES TAX",
            "Amendment of",
        ):
            idx = text.find(marker)
            if idx >= 0:
                start = idx
                break
        end = _find_block_end(text, start)
        body = text[start:end].strip()
        if len(body) > 80:
            return [("Goods and Services Tax Amendments", body)]

    part_starts = list(GST_PART_START.finditer(text))
    if part_starts:
        blocks: list[tuple[str, str]] = []
        for i, match in enumerate(part_starts):
            heading = match.group(0).replace("\n", " — ").strip()
            start = match.start()
            end = (
                part_starts[i + 1].start()
                if i + 1 < len(part_starts)
                else _find_block_end(text, start)
            )
            body = text[start:end].strip()
            if len(body) > 80:
                blocks.append((heading, body))
        return blocks

    sub_blocks = _extract_gst_subheading_blocks(text)
    if sub_blocks and sum(len(body) for _, body in sub_blocks) >= 1500:
        return sub_blocks

    section_body = _extract_gst_section_range(_enacted_text(text))
    if len(section_body) > 80:
        return [("Goods and Services Tax Amendments", section_body)]

    inline = GST_INLINE_START.search(text)
    if inline:
        start = inline.start()
        end = _find_block_end(text, start)
        body = text[start:end].strip()
        if len(body) > 80:
            return [("Goods and Services Tax Amendments", body)]

    return []


def split_into_sections(body: str) -> list[tuple[str, str, str]]:
    """Split part body into (section_no, section_title, section_text) tuples."""
    matches = list(SECTION_LINE.finditer(body))
    if not matches:
        return [("", "", body)]

    sections: list[tuple[str, str, str]] = []
    for i, m in enumerate(matches):
        sec_no = m.group(1)
        first_line = m.group(2).strip()
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(body)
        chunk = body[start:end].strip()
        title = first_line
        if len(first_line) > 120:
            title = first_line[:120] + "…"
        full = f"{sec_no}. {first_line}\n{chunk}".strip()
        sections.append((sec_no, title, full))
    return sections


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Arial"


def add_para(doc: Document, text: str, *, bold: bool = False, size: int = 11) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)


def add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def build_document() -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    add_heading(doc, "Goods and Services Tax", 0)
    add_heading(doc, "Finance Acts — GST Amendments (2017 onwards)", 1)
    add_para(
        doc,
        "Compiled from www.indiabudget.gov.in, gstcouncil.gov.in and egazette.gov.in. "
        f"Generated on {datetime.now().strftime('%d %B %Y, %H:%M')}.",
    )
    add_para(
        doc,
        "Disclaimer: Extracted for portal reference. Verify against official Gazette "
        "notifications and effective dates before relying on this text.",
    )

    add_heading(doc, "Contents", 1)
    for idx, act in enumerate(ACT_SOURCES, start=1):
        label = act["title"]
        if act.get("note_only"):
            add_para(doc, f"F{idx}. {label} (no GST amendments)")
        else:
            add_para(doc, f"F{idx}. {label}")

    add_page_break(doc)
    add_heading(doc, "Part F — Finance Acts", 1)

    for idx, act in enumerate(ACT_SOURCES, start=1):
        marker = f"F{idx}. {act['title']}"
        add_heading(doc, marker, 2)
        add_para(doc, f"Source: {act.get('source_label', 'Official publication')}", bold=True)
        add_para(doc, act.get("source_url", ""))
        if act.get("budget_url"):
            add_para(doc, f"Enacted text: {act['budget_url']}")

        if act.get("note_only"):
            add_para(doc, act["note"])
            continue

        pdf_name = act.get("pdf")
        pdf_path = os.path.join(SRC_DIR, pdf_name) if pdf_name else ""
        text = ""
        if pdf_path and os.path.isfile(pdf_path):
            text = pdf_text(pdf_path)
        blocks = extract_gst_blocks(text) if text else []

        if not blocks and act.get("fallback_pdf"):
            fb = os.path.join(SRC_DIR, act["fallback_pdf"])
            if os.path.isfile(fb):
                blocks = extract_gst_blocks(pdf_text(fb))

        if not blocks:
            add_para(doc, "[GST provisions could not be auto-extracted from the source PDF.]")
            continue

        for heading, body in blocks:
            add_heading(doc, heading.replace("\n", " "), 3)
            sections = split_into_sections(body)
            for sec_no, sec_title, sec_text in sections:
                if sec_no:
                    add_heading(doc, f"Section {sec_no} {sec_title}", 4)
                for para in sec_text.split("\n"):
                    line = para.strip()
                    if not line or NOISE_LINE.match(line):
                        continue
                    add_para(doc, line)

    return doc


def main() -> None:
    os.makedirs(SRC_DIR, exist_ok=True)
    doc = build_document()
    doc.save(OUT_DOCX)
    print(f"Wrote {OUT_DOCX}")


if __name__ == "__main__":
    main()